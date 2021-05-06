import os
import nmap
import datetime
import sender
import email, smtplib, ssl
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from docx.shared import Inches
from docx.shared import Pt
import subprocess
import re

#Print Logo 
logoRead = open('TextLogo.txt', 'r')
logoOutput = logoRead.read()
print(logoOutput)
logoRead.close()

#Start Document and set font
document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

#IP address array, discovered by Nmap
ipAddresses = []

#Regex to check email is valid
emailChecker = '^(\w|\.|\_|\-)+[@](\w|\_|\-|\.)+[.]\w{2,3}$'

#Create Report title page and how to use page
def reportCreation(document):
        document.add_picture('inquis.jpg', width=Inches(6), height=Inches(1.5))
        titlePage = document.add_paragraph("")
        titlePage.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titlePage.add_run("Smart Report from Penetration Test").font.size = Pt(26)
        lineBreak = titlePage.add_run()
        lineBreak.add_break()
        #date and time
        date = datetime.datetime.today()
        hour = date.hour
        minute = date.minute
        second = date.second
        day = date.day
        month = date.month
        year = date.year
        timeCheck = titlePage.add_run("Test conducted at: " + (str(hour) + ":" + str(minute) + ":" + str(second) + " on the " + str(day) + "/" + str(month) + "/" + str(year)))
        timeCheck.font.size = docx.shared.Pt(18)
       
        #Logo Placement
        addTitleImage = titlePage.add_run("")
        addTitleImage.add_picture('InquisitorLogo.png', width=Inches(5), height=Inches(6))
        document.add_page_break()

        #How to use Page
        document.add_heading('How to use', 0)
        document.add_paragraph('This Report provides a detailed analysis of your Network and/or Website domain. The Report has been structured into three distinc sections to match that of a standard Penetration Testing Model. These sections are as followed:')
        document.add_paragraph('Information Gathering: These stage aims to give an overview of the network/website, detailing the username, DNS records, and port activity through a series of scans. This is often the bread and butter of a Penetration test and provides information to be used in future stages.')
        document.add_paragraph('Vulnerability Scanning: The network and website are then scanned through a variety of tools, aiming to find any weaknesses. The vulnerabilities are the highlight of this report and need to be the key takeaway from reading.')
        document.add_paragraph('Exploitation Testing: This stage demonstrates the power a vulnerability can have over a network or website, and the devasting results it could have in the hands of a Malicious Hacker.')
        document.add_paragraph('All the findings have been stripped of irrelevant information and placed into this report, however the raw results from scans can be found in the text file attached to email. Both reports have been deleted off your system as to avoid someone using the contents to harm the network, rather then heal it.')
        document.add_page_break()
        
#whoami function with reporting
def whoamiTool():
        print("Conducting whoami command (1/5)")
        result = subprocess.Popen('whoami', stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout,_ = result.communicate()
        out = stdout.strip()
        out = out.strip()
        whoamiResult = out.decode('ascii')

        #open and add to raw file
        font.size = Pt(12)
        textFile = open("textdump.txt", "a")
        textFile.write("Result for whoami command: " + whoamiResult + "\n")
        textFile.write("----------------------------------------------------- \n")
        #write to smart report
        document.add_heading('Whomai Command', level=1)
        document.add_paragraph('Whoami is a simple linux command that displays the current users name')
        document.add_paragraph('Result for whoami command: %s' % whoamiResult)

#nslookup function with reporting        
def nslookupTool(domainName):
        print("Conducting nslookup command (2/5)")
        s = subprocess.Popen(['nslookup', domainName], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout,_ = s.communicate()
        out = stdout.strip()
        out = out.strip()
        nslookupResult = out.decode('ascii')
        #open and add to raw file
        font.size = Pt(12)
        textFile = open("textdump.txt", "a")
        textFile.write("Result for nslookup command: \n" + nslookupResult + "\n")
        textFile.write("----------------------------------------------------- \n")
        #write to smart report
        document.add_heading('NSLookup Command', level=1)
        document.add_paragraph('NSLookup is tool used to find more information about a domain name or other DNS records')
        document.add_paragraph('Result for nslookup command: \n %s' % nslookupResult)

#whois function with reporting 
def whoisTool(domainName):
        print("Conducting whois command (3/12)")
        strip, strippedName = domainName.split('www.')
        s = subprocess.Popen(['whois', strippedName], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout,_ = s.communicate()
        out = stdout.strip()
        out = out.strip()
        whoisRaw = out.decode('utf-8')
        try:
                whoisResult, stripped = whoisRaw.split('>>>', 1)
        except ValueError:
                whoisResult, stripped = whoisRaw.split('WHOIS lookup made at', 1)
                
        #open and add to raw file
        font.size = Pt(12)
        textFile = open("textdump.txt", "a")
        textFile.write("Result for whois command: \n" + whoisResult + "\n")
        textFile.write("----------------------------------------------------- \n")
        #write to smart report
        document.add_heading('Whois Command', level=1)
        document.add_paragraph('Whois is a listing system that show records about websites, such as ownership, domains, and other useful informaton')
        document.add_paragraph('Result for whois command: \n %s' % whoisResult)

def dmitryTool(domainName):
        print("Conducting DMitry scan (4/12)")
        strip, strippedName = domainName.split('www.')
        s = subprocess.Popen(['dmitry', '-nse', strippedName], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout,_ = s.communicate()
        out = stdout.strip()
        out = out.strip()
        dmitryRaw = out.decode('ascii')
        strip1, dmitryRaw2 = dmitryRaw.split('"There be some deep magic going on"')
        dmitryResult, strip2 = dmitryRaw2.split('All scans completed')
        #open and add to raw file
        font.size = Pt(12)
        textFile = open("textdump.txt", "a")
        textFile.write("Result for DMitry command: \n" + dmitryResult + "\n")
        textFile.write("----------------------------------------------------- \n")
        #write to smart report
        document.add_heading('DMitry Command', level=1)
        document.add_paragraph('DMitry (Deepmagic Information Gathering Tool) is a command line application used to find a wide variety of information about a host. DMitry is used here to find subdomains of the host and possible email addresses attached to it, whilst displaying basic Netcraft information.')
        document.add_paragraph('Result for whois command: \n %s' % dmitryResult)

def dnswalkTool(domainName):
        print("Conducting DNSWalki command (5/12)")
        strip, strippedName = domainName.split('www.')
        dnsName = strippedName+'.'
        s = subprocess.Popen(['dnswalk', dnsName], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout,_ = s.communicate()
        out = stdout.strip()
        out = out.strip()
        dnswalkRaw = out.decode('ascii')
        #open and add to raw file
        font.size = Pt(12)
        textFile = open("textdump.txt", "a")
        textFile.write("Result for DNSWalk command: \n" + dnswalkRaw + "\n")
        textFile.write("----------------------------------------------------- \n")
        #write to smart report
        document.add_heading('DNSWalk Command', level=1)
        document.add_paragraph('DNSWalk is a Domain Name System (DNS) debugger, checking a domain for consistency and accuracy through zone transfers.')
        document.add_paragraph('Result for DNSWalk command: \n %s' % dnswalkRaw)

#Nmap scan, discovery of hosts and reporting of results        
def nmapScanTool(ipaddr):
        print("Conducting Nmap scan (4/5)")
        font.size = Pt(12)
        document.add_heading('Nmap port scan', level=1)
        document.add_paragraph("Nmap is network scanning tool used in the majority of penetration tests, it can discover hosts to be used in further tests and services on the network, Nmap can also provide information on which ports are open.")
        #Start the Scan
        nmScan = nmap.PortScanner()
        nmScan.scan(ipaddr + ' -oN textdump.txt --append-output')
        #Output to the Smart Report
        for host in nmScan.all_hosts():
                document.add_paragraph('Host : %s (%s)' % (host, nmScan[host].hostname()))
                ipAddresses.append(host)
                document.add_paragraph('State : %s' % nmScan[host].state())
                for proto in nmScan[host].all_protocols():
                        document.add_paragraph('----------')
                        document.add_paragraph('Protocol : %s' % proto)
                        lport = nmScan[host][proto].keys()
                        for port in lport:
                                document.add_paragraph('port : %s\tstate : %s' % (port, nmScan[host][proto][port]['state']))
                document.add_paragraph('-------------------------------')

#OS Detection with discovered hosts and reporting       
def osDetectTool():
        print("Conducting OS detection (5/5)")
        amount = 0
        document.add_heading('OS Detection', level=1)
        document.add_paragraph('Nmap also provides an excellent OS Detection scan, using the result collected previously the entire networks Operating Systems can be found. This information can be useful for finding out more about unknown hosts or simply to map your network.')
        #Conducting an OS detection on each IP address found
        while amount < len(ipAddresses):
                try:
                        s = subprocess.Popen(['nmap', '-O', ipAddresses[amount], '-oN', 'textdump.txt', '--append-output'], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                        stdout,_ = s.communicate()
                        out = stdout.strip()
                        out = out.strip()
                        osDetectRaw = out.decode('ascii')
                        #Stripping repeated information
                        stripped, osStrip = osDetectRaw.split('Device type')
                        #Stripping Irrelevant Information
                        osDetectResult, stripped = osStrip.split('OS detection performed. Please report any incorrect results at')
                        #write to smart report
                        document.add_paragraph('Result for %s: ' % ipAddresses[amount])
                        document.add_paragraph('Device type %s' % osDetectResult)
                        document.add_paragraph('-------------------------------')
                #If no OS was detected        
                except ValueError:
                        document.add_paragraph("No OS was detected for %s" % ipAddresses[amount])
                        document.add_paragraph('-------------------------------')
                amount += 1        

                  
            
#Main with user inputs and cleanup        
def main():
        print("Enter the email of the tester (Reports will be sent to this address): ")
        validEmail = False
        #Checking if the Email address is valid
        while validEmail == False:
                reciever = str(input())
                if(re.search(emailChecker, reciever)):
                    validEmail = True
                else:
                    print("That is not a valid email: ")
        #Taking in the intial IP
        print("Enter a IP address to begin the Investigation: ")
        ipaddr = str(input())
        #Taking in the domain name
        print("Enter Domain name (i.e. www.name.com)")
        domainName = str(input())
        
        #Main operation of Inquisitor
        reportCreation(document)
        whoamiTool()
        nslookupTool(domainName)
        whoisTool(domainName)
        dmitryTool(domainName)
        
        nmapScanTool(ipaddr)
        osDetectTool()
        
        document.add_page_break()
        document.add_heading('Vulnerability Scanning')
        
        #Emailing and Cleanup
        document.save('report.docx')
        sender.emailSender(reciever)
        os.remove("report.docx")
        os.remove("textdump.txt")
        print("Report emailed and removed off device")



main()

